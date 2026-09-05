"""
docx_export.py
---------------
Turns a validated MeetingExtraction into a polished .docx file:
- Title + summary
- Key decisions with confidence tags
- Action items grouped by owner (unassigned last)
- Ambiguities/flags section (from validation.py) so reviewers see it too
"""

from __future__ import annotations

from io import BytesIO
from typing import Dict, List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

from schema import ActionItem, MeetingExtraction
from validation import ValidationReport

CONFIDENCE_COLORS = {
    "high": RGBColor(0x1E, 0x7B, 0x34),
    "medium": RGBColor(0xB8, 0x86, 0x0B),
    "low": RGBColor(0xB0, 0x30, 0x30),
}


def _set_cell_shading(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.makeelement(qn("w:shd"), {qn("w:fill"): hex_color})
    tc_pr.append(shd)


def _add_heading(doc: Document, text: str, level: int = 1):
    heading = doc.add_heading(text, level=level)
    return heading


def _add_confidence_run(paragraph, confidence: str):
    run = paragraph.add_run(f"  [{confidence.upper()} CONFIDENCE]")
    run.font.size = Pt(9)
    run.bold = True
    run.font.color.rgb = CONFIDENCE_COLORS.get(confidence, RGBColor(0x44, 0x44, 0x44))


def _group_by_owner(items: List[ActionItem]) -> Dict[str, List[ActionItem]]:
    groups: Dict[str, List[ActionItem]] = {}
    for item in items:
        key = item.owner if item.owner else "Unassigned"
        groups.setdefault(key, []).append(item)
    # Unassigned last, everyone else alphabetical
    ordered_keys = sorted(k for k in groups if k != "Unassigned")
    if "Unassigned" in groups:
        ordered_keys.append("Unassigned")
    return {k: groups[k] for k in ordered_keys}


def build_docx(
    extraction: MeetingExtraction,
    validation_report: Optional[ValidationReport] = None,
    source_filename: Optional[str] = None,
) -> BytesIO:
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading(extraction.meeting_title, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    if source_filename:
        note = doc.add_paragraph()
        run = note.add_run(f"Source: {source_filename}")
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    _add_heading(doc, "Summary", level=1)
    doc.add_paragraph(extraction.summary)

    _add_heading(doc, "Key Decisions", level=1)
    if extraction.key_decisions:
        for decision in extraction.key_decisions:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(decision.decision)
            _add_confidence_run(p, decision.confidence)
    else:
        doc.add_paragraph("No firm decisions were recorded in this meeting.").italic = True

    _add_heading(doc, "Action Items (Grouped by Owner)", level=1)
    if extraction.action_items:
        grouped = _group_by_owner(extraction.action_items)
        for owner, items in grouped.items():
            _add_heading(doc, owner, level=2)
            table = doc.add_table(rows=1, cols=3)
            table.style = "Light Grid Accent 1"
            hdr = table.rows[0].cells
            hdr[0].text = "Task"
            hdr[1].text = "Deadline"
            hdr[2].text = "Confidence / Evidence"
            for item in items:
                row = table.add_row().cells
                row[0].text = item.task
                row[1].text = item.deadline if item.deadline else "Not stated"
                conf_cell = row[2]
                conf_cell.text = f"{item.confidence.upper()}"
                ev_para = conf_cell.add_paragraph()
                ev_run = ev_para.add_run(f'"{item.evidence}"')
                ev_run.italic = True
                ev_run.font.size = Pt(9)
            doc.add_paragraph()
    else:
        doc.add_paragraph("No action items were extracted from this transcript.").italic = True

    _add_heading(doc, "Ambiguities & Open Questions", level=1)
    if extraction.ambiguities:
        for amb in extraction.ambiguities:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(f"{amb.issue}")
            run.bold = True
            p.add_run(f"  (related to: {amb.related_item})")
    else:
        doc.add_paragraph("None identified by the model.").italic = True

    if validation_report and validation_report.has_issues:
        _add_heading(doc, "Automated Validation Flags", level=1)
        note = doc.add_paragraph()
        note.add_run(
            "These items were flagged by a rule-based check that compares each item's "
            "evidence quote against the raw transcript. They are not necessarily wrong — "
            "review before treating them as confirmed."
        ).italic = True
        for flagged in validation_report.flagged_action_items:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(f"{flagged.item.task}").bold = True
            for reason in flagged.reasons:
                sub = doc.add_paragraph(style="List Bullet 2")
                sub.add_run(reason).font.size = Pt(9)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
